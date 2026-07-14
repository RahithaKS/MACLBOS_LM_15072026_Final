"""
Document processing API routes
"""
from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from typing import List, Optional
import os
import sys
import logging

# Add parent directories to path to import your existing modules
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../..'))

from database import update_processing_status, get_processing_status
from config import settings

logger = logging.getLogger(__name__)

router = APIRouter()

# Try to import your existing modules
try:
    from services.document_processor import process_single_document as process_document, get_file_format, validate_file_format
    from services.company_extractor import CompanyExtractor
    HAS_PROCESSORS = True
except ImportError as e:
    logger.warning(f"Could not import document processors: {e}")
    HAS_PROCESSORS = False

from pydantic import BaseModel

class ProcessDocumentRequest(BaseModel):
    file_path: str
    ai_config: Optional[dict] = None

@router.post("/process/{document_id}")
async def process_document_endpoint(
    document_id: str,
    request: ProcessDocumentRequest,
    background_tasks: BackgroundTasks
):
    """
    Process a document: extract text, chunk, generate embeddings
    """
    if not HAS_PROCESSORS:
        raise HTTPException(status_code=501, detail="Document processors not available")
    
    try:
        # Update status to processing
        update_processing_status(document_id, 'processing')
        
        # Schedule background task
        background_tasks.add_task(
            process_document_background,
            document_id,
            request.file_path,
            request.ai_config,
        )
        
        return {
            "success": True,
            "message": "Document processing started",
            "document_id": document_id
        }
        
    except Exception as e:
        logger.error(f"Failed to start document processing: {e}")
        update_processing_status(document_id, 'failed', error_message=str(e))
        raise HTTPException(status_code=500, detail=str(e))

async def process_document_background(document_id: str, file_path: str, ai_config: dict = None):
    """Background task for document processing"""
    try:
        import sys
        import os
        sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../..'))
        
        from services.document_processor import process_single_document as process_document
        from services.company_extractor import CompanyExtractor
        from services.vector_store import chunk_text, get_embeddings
        from database import store_document_chunks, store_vector_embeddings
        from config import get_default_ai_config

        # Use passed ai_config, or fall back to global Azure env vars, or Ollama
        effective_ai_config = ai_config or get_default_ai_config()
        
        logger.info(f"Processing document {document_id}")
        
        # Convert relative path to absolute path
        # Python backend runs from python_backend/, but uploads are in project root
        if not os.path.isabs(file_path):
            # Get project root (two levels up from python_backend/api/routes/)
            project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../..'))
            
            # If file_path doesn't start with 'uploads/', prepend it
            if not file_path.startswith('uploads/'):
                file_path = os.path.join('uploads', file_path)
            
            absolute_path = os.path.join(project_root, file_path)
        else:
            absolute_path = file_path
        
        logger.info(f"Using absolute path: {absolute_path}")
        
        # Verify file exists
        if not os.path.exists(absolute_path):
            raise FileNotFoundError(f"File not found: {absolute_path}")
        
        # Read file content and get filename
        with open(absolute_path, 'rb') as f:
            file_content = f.read()
        filename = os.path.basename(absolute_path)
        
        # Process document (expects bytes content and filename)
        processed_data = process_document(file_content, filename)
        
        # Extract company name
        extractor = CompanyExtractor()
        filename = os.path.basename(file_path)
        
        # Combine all text content
        full_text = ""
        if isinstance(processed_data, dict) and 'content' in processed_data:
            content = processed_data['content']
            if isinstance(content, list):
                full_text = '\n\n'.join([page.get('text', '') for page in content if isinstance(page, dict)])
            elif isinstance(content, str):
                full_text = content
        elif isinstance(processed_data, list):
            full_text = '\n\n'.join([item.get('text', '') for item in processed_data if isinstance(item, dict)])
        
        company_info = extractor.extract_company_name(full_text, filename, ai_config=effective_ai_config)
        company_name = company_info.get('primary_company')
        
        # Chunk the text
        chunks = chunk_text(full_text, doc_id=document_id)
        
        # Store chunks in database
        chunk_ids = store_document_chunks(document_id, chunks)
        
        # Update processing status with chunk count
        update_processing_status(
            document_id,
            'processing',
            total_chunks=len(chunks),
            processed_chunks=0,
            company_name=company_name
        )
        
        # Generate and store embeddings (Azure 3072-dim or Ollama 1024-dim depending on config)
        texts = [chunk['text'] for chunk in chunks]
        embeddings, _ = get_embeddings(texts, ai_config=effective_ai_config)
        use_azure = (
            effective_ai_config is not None
            and effective_ai_config.get("provider") == "azure_openai"
            and effective_ai_config.get("embedding_model")
        )
        store_vector_embeddings(chunk_ids, embeddings, use_azure=use_azure)
        
        # Mark as completed with final status
        update_processing_status(
            document_id,
            'completed',
            total_chunks=len(chunks),
            processed_chunks=len(chunks)
        )
        
        logger.info(f"Document {document_id} processed successfully")
        
    except Exception as e:
        logger.error(f"Document processing failed: {e}")
        update_processing_status(
            document_id,
            'failed',
            error_message=str(e)
        )

@router.get("/{document_id}/status")
async def get_document_status(document_id: str):
    """Get document processing status"""
    try:
        status = get_processing_status(document_id)
        
        if not status:
            return {
                "document_id": document_id,
                "status": "not_found",
                "message": "No processing status found"
            }
        
        return status
        
    except Exception as e:
        logger.error(f"Failed to get document status: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/formats")
async def get_supported_formats():
    """Get list of supported document formats"""
    return {
        "formats": [
            {"ext": "pdf", "name": "PDF Document", "mime": "application/pdf"},
            {"ext": "docx", "name": "Word Document", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
            {"ext": "xlsx", "name": "Excel Spreadsheet", "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
            {"ext": "csv", "name": "CSV File", "mime": "text/csv"},
            {"ext": "txt", "name": "Text File", "mime": "text/plain"},
            {"ext": "json", "name": "JSON File", "mime": "application/json"},
            {"ext": "xml", "name": "XML File", "mime": "application/xml"},
            {"ext": "html", "name": "HTML File", "mime": "text/html"},
            {"ext": "png", "name": "PNG Image", "mime": "image/png"},
            {"ext": "jpg", "name": "JPEG Image", "mime": "image/jpeg"},
        ]
    }


class ExtractTextRequest(BaseModel):
    file_path: str


@router.post("/extract-text")
async def extract_text_from_file(request: ExtractTextRequest):
    """
    Extract text from a document file.
    Supports: .txt, .docx, .doc, .pdf, .md, .csv, .xlsx
    Security: Only allows files within the uploads directory.
    """
    import os
    
    file_path = request.file_path
    
    # Security: Normalize path and prevent path traversal
    uploads_base = os.path.abspath(os.path.join(os.getcwd(), 'uploads'))
    normalized_path = os.path.abspath(file_path)
    
    # Ensure file path is within allowed directories
    if not normalized_path.startswith(uploads_base):
        # Also allow attached_assets for development/testing
        attached_base = os.path.abspath(os.path.join(os.getcwd(), 'attached_assets'))
        if not normalized_path.startswith(attached_base):
            logger.warning(f"[SECURITY] Path traversal attempt blocked: {file_path}")
            raise HTTPException(status_code=403, detail="Access denied")
    
    if not os.path.exists(normalized_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    file_path = normalized_path
    
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    try:
        text = ""
        
        if ext in ['.txt', '.md', '.csv']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        elif ext in ['.doc', '.docx']:
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            paragraphs.append(row_text)
                text = '\n\n'.join(paragraphs)
            except Exception as e:
                logger.error(f"Failed to extract text from Word document: {e}")
                raise HTTPException(status_code=500, detail="Failed to parse Word document")
        
        elif ext == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    pages = []
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages.append(page_text)
                    text = '\n\n'.join(pages)
            except Exception as e:
                logger.error(f"Failed to extract text from PDF: {e}")
                raise HTTPException(status_code=500, detail="Failed to parse PDF document")
        
        elif ext == '.xlsx':
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True)
                sheets = []
                for sheet in wb.worksheets:
                    rows = []
                    for row in sheet.iter_rows(values_only=True):
                        row_vals = [str(cell) if cell is not None else '' for cell in row]
                        if any(v for v in row_vals):
                            rows.append(' | '.join(row_vals))
                    if rows:
                        sheets.append(f"Sheet: {sheet.title}\n" + '\n'.join(rows))
                text = '\n\n'.join(sheets)
                wb.close()
            except Exception as e:
                logger.error(f"Failed to extract text from Excel: {e}")
                raise HTTPException(status_code=500, detail="Failed to parse Excel spreadsheet")
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {ext}")
        
        return {
            "success": True,
            "text": text,
            "format": ext
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Text extraction failed for {ext} file: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract text from document")
